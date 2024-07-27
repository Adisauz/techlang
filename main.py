import streamlit as st
from googletrans import Translator, LANGUAGES
from docx import Document as DocxDocument
import fitz  # PyMuPDF
import os

def translate_text(text, target_language):
    translator = Translator() 
    try:
        translation = translator.translate(text, dest=target_language)
        return translation.text
    except Exception as e:
        st.error(f"Translation error: {e}")
        return ""

def create_word_doc(text, filename):
    doc = DocxDocument()
    doc.add_paragraph(text)
    doc.save(filename)

def read_pdf_text(pdf_file):
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def read_docx_text(docx_file):
    doc = DocxDocument(docx_file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def read_txt_text(txt_file):
    return txt_file.read().decode('utf-8')

st.title("Educational Content Translator")

uploaded_file = st.file_uploader("Upload a DOCX, PDF, or TXT document", type=["pdf", "docx", "txt"])
input_text = ""

if uploaded_file:
    file_type = uploaded_file.type
    if file_type == "application/pdf":
        input_text = read_pdf_text(uploaded_file)
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        input_text = read_docx_text(uploaded_file)
    elif file_type == "text/plain":
        input_text = read_txt_text(uploaded_file)
    else:
        st.error("Unsupported file type!")

if input_text:
    language_options = list(LANGUAGES.values())
    language_keys = list(LANGUAGES.keys())
    selected_language = st.selectbox("Select the target language", options=language_options)
    target_language = language_keys[language_options.index(selected_language)]
    output_filename = st.text_input("Enter the output Word filename (e.g., 'translated.docx'):")

    if st.button("Translate and Generate Word Document"):
        if target_language and output_filename:
            translated_text = translate_text(input_text, target_language)
            
            # Get the path to the Downloads folder
            downloads_path = os.path.join(os.path.expanduser("~"), "Downloads")
            save_path = os.path.join(downloads_path, output_filename)
            
            create_word_doc(translated_text, save_path)
            st.success(f"Translation completed and saved as {save_path}")
        else:
            st.error("Please fill in all the fields.")
else:
    st.info("Please upload a document to translate.")
